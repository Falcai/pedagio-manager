import tkinter as tk
from tkinter import filedialog, ttk
import json
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.enum.table import WD_ALIGN_VERTICAL
from datetime import datetime

# Constantes
JSON_FILE = "pedagios.json"

# Dicionário para armazenar os pedágios com seus dados
pedagios_data = {}

# Funções de manipulação de arquivos e dados
def open_doc(file_path):
    """Abre um documento Word."""
    try:
        doc = Document(file_path)
    except Exception as e:
        message_var.set(f"Erro: Não foi possível abrir o arquivo: {e}")
        return None
    return doc

def add_entry(doc, data, local, valor, cupom_fiscal, file_path):
    """Adiciona uma entrada na tabela do documento."""
    table = doc.tables[0]
    num_ordem = len(table.rows)
    row_cells = table.add_row().cells

    # Preenche as células da nova linha
    row_cells[0].text = str(num_ordem)  # Nº Ordem
    row_cells[1].text = data  # Data
    row_cells[2].text = f"R${float(valor):.2f}".replace('.', ',')  # Valor
    row_cells[3].text = local  # Local

    # Documento Fiscal - Utilização
    pedagio_info = pedagios_data.get(local)
    if pedagio_info:
        doc_fiscal = (f"Cupom Fiscal nº\n{cupom_fiscal}\nCNPJ Nº {pedagio_info[0]['cnpj']}\n"
                      f"Razão Social: {pedagio_info[0]['razao_social']}\n")
        run = row_cells[4].paragraphs[0].add_run(doc_fiscal)
        run = row_cells[4].paragraphs[0].add_run("- Pedágio")
        run.font.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
    else:
        message_var.set("Erro: Pedágio não encontrado.")
        return

    format_row_cells(row_cells)
    add_borders_to_cells(row_cells)
    doc.save(file_path)

    # Atualiza a mensagem e o contador de linhas
    message_var.set("Entrada adicionada com sucesso!")
    lines_added_var.set(f"Total de linhas adicionadas: {num_ordem}")

def format_row_cells(row_cells):
    """Formata as células da linha adicionada."""
    for cell in row_cells:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
        if not cell.paragraphs:
            cell.add_paragraph()
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            run.font.name = 'Calibri'
            run.font.size = Pt(11)

def add_borders_to_cells(cells):
    """Adiciona bordas às células."""
    for cell in cells:
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ["top", "left", "bottom", "right"]:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tcBorders.append(border)
        tcPr.append(tcBorders)

def format_date(date_str):
    """Formata a data para o formato dd/mm/aaaa."""
    if len(date_str) == 8:
        return f"{date_str[:2]}/{date_str[2:4]}/{date_str[4:]}"
    else:
        raise ValueError("Erro: Data inválida. Use o formato ddmmaaaa sem barras.")

def submit(*args):
    """Processa a submissão do formulário."""
    data = entry_data.get()
    valor = entry_valor.get().replace(',', '.')  # Converte para float
    cupom_fiscal = entry_cupom_fiscal.get()
    file_path = file_path_entry.get()

    # Verificação de campos obrigatórios
    if not file_path:
        message_var.set("Erro: Por favor, selecione o caminho do documento.")
        return

    if not data or not valor or not cupom_fiscal:
        message_var.set("Erro: Todos os campos devem ser preenchidos.")
        return

    try:
        formatted_date = format_date(data)
        datetime.strptime(formatted_date, '%d/%m/%Y')
    except ValueError as ve:
        message_var.set(str(ve))
        return

    doc = open_doc(file_path)
    if doc:
        local = next((key for key, value in pedagios_data.items() if any(d['valor'] == valor for d in value)), None)
        if local:
            add_entry(doc, formatted_date, local, float(valor), cupom_fiscal, file_path)
        else:
            message_var.set("Erro: Pedágio não encontrado.")

def browse_file():
    """Abre um diálogo para selecionar um arquivo."""
    file_path = filedialog.askopenfilename(filetypes=[("Documentos Word", "*.docx")])
    if file_path:
        file_path_entry.delete(0, tk.END)
        file_path_entry.insert(0, file_path)

def open_pedagio_manager():
    """Abre a janela de gerenciamento de pedágios."""
    manager_window = tk.Toplevel(root)
    manager_window.title("Gerenciador de Pedágios")
    
    def update_pedagios_list():
        """Atualiza a lista de pedágios na interface."""
        for i in tree.get_children():
            tree.delete(i)
        for key, value in pedagios_data.items():
            for pedagio in value:
                tree.insert('', 'end', values=(key, pedagio['cnpj'], pedagio['razao_social'], f"R${pedagio['valor']}"))

    def add_pedagio():
        """Adiciona um novo pedágio."""
        local = entry_local.get()
        cnpj = entry_cnpj.get()
        razao_social = entry_razao_social.get()
        valor = entry_valor_pedagio.get().replace(',', '.')  # Converte para float
        if local and cnpj and razao_social and valor:
            if local in pedagios_data:
                pedagios_data[local].append({"cnpj": cnpj, "razao_social": razao_social, "valor": valor})
            else:
                pedagios_data[local] = [{"cnpj": cnpj, "razao_social": razao_social, "valor": valor}]
            save_pedagios_data()
            update_pedagios_list()
            message_var.set("Pedágio adicionado com sucesso!")
        else:
            message_var.set("Erro: Todos os campos são obrigatórios para adicionar um pedágio.")
    
    def delete_pedagio():
        """Remove um pedágio selecionado."""
        selected_item = tree.selection()
        if selected_item:
            item = tree.item(selected_item)
            local = item['values'][0]
            cnpj = item['values'][1]
            pedagios_data[local] = [pedagio for pedagio in pedagios_data[local] if pedagio['cnpj'] != cnpj]
            if not pedagios_data[local]:
                del pedagios_data[local]
            save_pedagios_data()
            update_pedagios_list()
            message_var.set("Pedágio removido com sucesso!")
        else:
            message_var.set("Erro: Selecione um pedágio para remover.")
    
    def update_pedagio():
        """Atualiza um pedágio selecionado."""
        selected_item = tree.selection()
        if selected_item:
            item = tree.item(selected_item)
            local = item['values'][0]
            cnpj = item['values'][1]
            new_cnpj = entry_cnpj.get()
            new_razao_social = entry_razao_social.get()
            new_valor = entry_valor_pedagio.get().replace(',', '.')  # Converte para float
            if new_cnpj and new_razao_social and new_valor:
                pedagio_index = next(index for index, pedagio in enumerate(pedagios_data[local]) if pedagio['cnpj'] == cnpj)
                pedagios_data[local][pedagio_index] = {"cnpj": new_cnpj, "razao_social": new_razao_social, "valor": new_valor}
                save_pedagios_data()
                update_pedagios_list()
                message_var.set("Pedágio atualizado com sucesso!")
            else:
                message_var.set("Erro: Todos os campos são obrigatórios para atualizar um pedágio.")
        else:
            message_var.set("Erro: Selecione um pedágio para atualizar.")

    # Interface do gerenciador de pedágios
    tk.Label(manager_window, text="Local:").grid(row=0, column=0, padx=10, pady=5)
    entry_local = tk.Entry(manager_window)
    entry_local.grid(row=0, column=1, padx=10, pady=5)

    tk.Label(manager_window, text="CNPJ:").grid(row=1, column=0, padx=10, pady=5)
    entry_cnpj = tk.Entry(manager_window)
    entry_cnpj.grid(row=1, column=1, padx=10, pady=5)

    tk.Label(manager_window, text="Razão Social:").grid(row=2, column=0, padx=10, pady=5)
    entry_razao_social = tk.Entry(manager_window)
    entry_razao_social.grid(row=2, column=1, padx=10, pady=5)

    tk.Label(manager_window, text="Valor em R$ (R$xx,xx):").grid(row=3, column=0, padx=10, pady=5)
    entry_valor_pedagio = tk.Entry(manager_window)
    entry_valor_pedagio.grid(row=3, column=1, padx=10, pady=5)

    tk.Button(manager_window, text="Adicionar", command=add_pedagio).grid(row=4, column=0, padx=10, pady=5)
    tk.Button(manager_window, text="Remover", command=delete_pedagio).grid(row=4, column=1, padx=10, pady=5)
    tk.Button(manager_window, text="Atualizar", command=update_pedagio).grid(row=4, column=2, padx=10, pady=5)

    columns = ('local', 'cnpj', 'razao_social', 'valor')
    tree = ttk.Treeview(manager_window, columns=columns, show='headings')
    tree.heading('local', text='Local')
    tree.heading('cnpj', text='CNPJ')
    tree.heading('razao_social', text='Razão Social')
    tree.heading('valor', text='Valor')

    tree.grid(row=5, column=0, columnspan=3, padx=10, pady=10)

    update_pedagios_list()

def save_pedagios_data():
    """Salva os dados dos pedágios em um arquivo JSON."""
    with open(JSON_FILE, 'w') as file:
        json.dump(pedagios_data, file)

def load_pedagios_data():
    """Carrega os dados dos pedágios de um arquivo JSON."""
    try:
        with open(JSON_FILE, 'r') as file:
            return json.load(file)
    except FileNotFoundError:
        return {}

# Carrega os pedágios já cadastrados
pedagios_data = load_pedagios_data()

# Configuração da interface gráfica principal
root = tk.Tk()
root.title("Automatização de Preenchimento de Tabela")

def format_data_entry(event=None):
    """Formata a entrada de data."""
    data = entry_data.get().replace("/", "")
    if len(data) == 10 and data[2] == data[5] == "/":
        formatted_data = f"{data[:2]}{data[3:5]}{data[6:]}"
        entry_data.delete(0, tk.END)
        entry_data.insert(0, formatted_data)

tk.Label(root, text="Caminho do Documento:").grid(row=0, column=0, padx=10, pady=5)
file_path_entry = tk.Entry(root, width=50)
file_path_entry.grid(row=0, column=1, padx=10, pady=5)
tk.Button(root, text="Procurar", command=browse_file).grid(row=0, column=2, padx=10, pady=5)

tk.Label(root, text="Data (ddmmaaaa):").grid(row=1, column=0, padx=10, pady=5)
entry_data = tk.Entry(root)
entry_data.grid(row=1, column=1, padx=10, pady=5)
entry_data.bind('<FocusOut>', format_data_entry)

tk.Label(root, text="Valor em R$ (R$xx,xx):").grid(row=2, column=0, padx=10, pady=5)
entry_valor = tk.Entry(root)
entry_valor.grid(row=2, column=1, padx=10, pady=5)

tk.Label(root, text="Número do Cupom Fiscal:").grid(row=3, column=0, padx=10, pady=5)
entry_cupom_fiscal = tk.Entry(root)
entry_cupom_fiscal.grid(row=3, column=1, padx=10, pady=5)

submit_button = tk.Button(root, text="Adicionar Entrada", command=submit)
submit_button.grid(row=4, column=0, columnspan=3, pady=10)

message_var = tk.StringVar()
lines_added_var = tk.StringVar()
message_label = tk.Label(root, textvariable=message_var)
message_label.grid(row=5, column=0, columnspan=3)
lines_added_label = tk.Label(root, textvariable=lines_added_var)
lines_added_label.grid(row=6, column=0, columnspan=3, pady=5)

tk.Button(root, text="Gerenciar Pedágios", command=open_pedagio_manager).grid(row=7, column=0, columnspan=3, pady=10)

# Inicializa os rótulos vazios
message_var.set("")
lines_added_var.set("Total de linhas adicionadas: 0")

root.bind('<Return>', submit)
root.bind_class("Entry", "<Tab>", lambda e: e.widget.tk_focusNext().focus())
root.bind_class("Entry", "<Shift-Tab>", lambda e: e.widget.tk_focusPrev().focus())

# Ordena o foco na sequência desejada
entry_data.focus()

root.mainloop()
